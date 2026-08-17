import os
import json
import time
import random
from flask import Flask, request, jsonify
import psycopg
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from groq import Groq
import warnings

warnings.filterwarnings("ignore", category=UserWarning, module="huggingface_hub.*")
# Load .env from the root directory
load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env'))
DATABASE_URL = os.getenv("DATABASE_URL")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

app = Flask(__name__)

print("Loading local embedding model (all-MiniLM-L6-v2)...")
# Note: we use all-MiniLM-L6-v2 since our db is vector(384)
model = SentenceTransformer('all-MiniLM-L6-v2')
print("Model loaded successfully.")

# Initialize Groq client
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

TRACK_RUBRICS = {}
DEFAULT_DIMENSIONS = ["correctness", "depth", "communication", "problem_solving"]

def load_track_rubrics():
    global TRACK_RUBRICS
    try:
        with psycopg.connect(DATABASE_URL) as conn:
            with conn.cursor() as cur:
                # We check if the table exists first (in case the migration hasn't run yet)
                cur.execute("""
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables 
                        WHERE table_name = 'track_rubrics'
                    );
                """)
                if cur.fetchone()[0]:
                    cur.execute("SELECT role_track, round_type, dimension, weight, description FROM track_rubrics")
                    rows = cur.fetchall()
                    new_rubrics = {}
                    for track, round_type, dim, weight, desc in rows:
                        # Use empty string if round_type is None so we can reliably key by a tuple
                        rt_key = round_type if round_type else ""
                        key = (track, rt_key)
                        if key not in new_rubrics:
                            new_rubrics[key] = []
                        new_rubrics[key].append({
                            "dimension": dim,
                            "weight": float(weight),
                            "description": desc
                        })
                    TRACK_RUBRICS = new_rubrics
                    print(f"Loaded rubrics for {len(TRACK_RUBRICS)} track configurations.")
                else:
                    print("Table track_rubrics does not exist yet. Using default dimensions.")
    except Exception as e:
        print(f"Error loading track rubrics: {e}")

load_track_rubrics()

@app.post("/admin/reload-rubrics")
def reload_rubrics():
    load_track_rubrics()
    return jsonify({"status": "ok", "count": len(TRACK_RUBRICS)})

@app.post("/admin/rubrics")
def write_rubric():
    """
    Accepts:
    {
      "role_track": "coding_dsa",
      "round_type": null,
      "dimensions": [
        {"dimension": "correctness", "weight": 0.35, "description": "..."},
        ...
      ]
    }
    """
    data = request.json
    role_track = data.get("role_track")
    round_type = data.get("round_type")
    dimensions = data.get("dimensions", [])
    
    if not role_track or not dimensions:
        return jsonify({"error": "role_track and dimensions are required"}), 400
        
    # Validate weights sum to 1.0
    total_weight = sum(d.get("weight", 0) for d in dimensions)
    if abs(total_weight - 1.0) > 0.001:
        return jsonify({
            "error": f"Weights for ({role_track}, {round_type}) must sum to 1.0. Actual sum: {total_weight}"
        }), 400
        
    with psycopg.connect(DATABASE_URL) as conn:
        with conn.cursor() as cur:
            # Upsert the dimensions
            for d in dimensions:
                # Using PostgreSQL ON CONFLICT requires a unique constraint/index.
                # Assuming (role_track, round_type, dimension) is unique per the migration.
                cur.execute("""
                    INSERT INTO track_rubrics (role_track, round_type, dimension, weight, description)
                    VALUES (%s, %s, %s, %s, %s)
                    ON CONFLICT (role_track, round_type, dimension) 
                    DO UPDATE SET weight = EXCLUDED.weight, description = EXCLUDED.description
                """, (role_track, round_type, d["dimension"], d["weight"], d.get("description", "")))
                
    load_track_rubrics()
    return jsonify({"status": "ok", "message": f"Saved {len(dimensions)} dimensions for {role_track}"})

@app.get("/health")
def health():
    return {"status": "ok", "model_loaded": model is not None}

@app.post("/next-question")
def next_question():
    """
    Accepts:
    {
      "target_gap": "optional area to dig into",
      "role_track": "frontend",
      "difficulty": "medium",
      "asked_question_ids": ["uuid1", "uuid2"],
      "session_id": "uuid"
    }
    """
    data = request.json
    target_gap = data.get("target_gap", "")
    role_track = data.get("role_track", "frontend")
    difficulty = data.get("difficulty", "medium")
    asked_question_ids = data.get("asked_question_ids", [])
    session_id = data.get("session_id")
    
    start_time = time.time()
    row = None
    has_tailored_context = False
    
    if session_id:
        try:
            with psycopg.connect(DATABASE_URL) as conn:
                with conn.cursor() as cur:
                    # Check if session has both resume_id and job_description_id, and they have non-null embeddings
                    cur.execute("""
                        SELECT 
                            s.resume_id, 
                            s.job_description_id,
                            r.embedding IS NOT NULL as has_r,
                            jd.embedding IS NOT NULL as has_jd
                        FROM interview_sessions s
                        LEFT JOIN resumes r ON s.resume_id = r.id
                        LEFT JOIN job_descriptions jd ON s.job_description_id = jd.id
                        WHERE s.id = %s
                    """, (session_id,))
                    s_row = cur.fetchone()
                    if s_row and s_row[0] and s_row[1] and s_row[2] and s_row[3]:
                        has_tailored_context = True
                        
                    if has_tailored_context:
                        # Select top 10 most relevant questions by average cosine similarity
                        # (since <=> is cosine distance, smaller average distance = higher similarity)
                        query = """
                            SELECT q.id, q.text, q.role_track, q.topic, q.difficulty, q.ideal_answer_points
                            FROM questions q
                            JOIN interview_sessions s ON s.id = %s
                            JOIN resumes r ON s.resume_id = r.id
                            JOIN job_descriptions jd ON s.job_description_id = jd.id
                            WHERE q.role_track = %s AND q.difficulty = %s AND q.id != ALL(%s)
                              AND r.embedding IS NOT NULL AND jd.embedding IS NOT NULL
                            ORDER BY ( (q.embedding <=> r.embedding) + (q.embedding <=> jd.embedding) ) / 2.0 ASC
                            LIMIT 10;
                        """
                        cur.execute(query, (session_id, role_track, difficulty, asked_question_ids))
                        rows = cur.fetchall()
                        if rows:
                            row = random.choice(rows)
                            print(f"Tailored question selection selected from {len(rows)} candidate questions.")
        except Exception as e:
            print(f"Error executing tailored question selection: {e}")
            
    if not row:
        # Fall back to existing selection logic
        search_text = target_gap if target_gap else f"{role_track} {difficulty} interview question"
        embedding = model.encode(search_text).tolist()
        
        try:
            with psycopg.connect(DATABASE_URL) as conn:
                with conn.cursor() as cur:
                    query = """
                        SELECT id, text, role_track, topic, difficulty, ideal_answer_points
                        FROM questions
                        WHERE role_track = %s AND difficulty = %s AND id != ALL(%s)
                        ORDER BY embedding <=> %s::vector
                        LIMIT 1;
                    """
                    cur.execute(query, (role_track, difficulty, asked_question_ids, embedding))
                    row = cur.fetchone()
        except Exception as e:
            print(f"Error executing fallback question selection: {e}")
            
    elapsed_time = time.time() - start_time
    print(f"Question selection took {elapsed_time:.3f}s (tailored: {has_tailored_context})")
    
    if not row:
        return jsonify({"error": "No more questions available"}), 404
        
    return jsonify({
        "id": row[0],
        "text": row[1],
        "role_track": row[2],
        "topic": row[3],
        "difficulty": row[4],
        "ideal_answer_points": row[5],
        "latency_sec": elapsed_time
    })

@app.post("/grade")
def grade():
    """
    Holistic transcript grading.
    Accepts:
    {
       "session_id": "uuid"
    }
    """
    if not client:
        return jsonify({"error": "GROQ_API_KEY not set"}), 500
        
    data = request.json
    session_id = data.get("session_id")
    
    with psycopg.connect(DATABASE_URL) as conn:
        with conn.cursor() as cur:
            # Fetch candidate name, track, and round type
            cur.execute("SELECT candidate_name, role_track, round_type FROM interview_sessions WHERE id = %s", (session_id,))
            session_row = cur.fetchone()
            candidate_name = session_row[0] if session_row and session_row[0] else "the candidate"
            role_track = session_row[1] if session_row else ""
            round_type = session_row[2] if session_row and session_row[2] else ""

            # Fetch the entire transcript for this session
            cur.execute("""
                SELECT speaker, content 
                FROM turns 
                WHERE session_id = %s 
                ORDER BY sequence_number ASC
            """, (session_id,))
            rows = cur.fetchall()
            
            if not rows:
                return jsonify({"error": "No transcript found for session"}), 404
            
            transcript = "\n\n".join([f"{r[0].upper()}:\n{r[1]}" for r in rows])
            
            # NOTE: rubric selection is session-level. A single interview_sessions.round_type 
            # determines which rubric weights apply to the whole session's Scorecard, not per-question.
            # If a session needs to mix round types, that requires per-turn rubric application instead.
            
            # Lookup specific dimensions for this track + round_type
            key = (role_track, round_type)
            # Fall back to base role track if specific round_type not configured
            fallback_key = (role_track, "")
            
            dims_info = TRACK_RUBRICS.get(key) or TRACK_RUBRICS.get(fallback_key)
            if dims_info:
                dimensions = [d["dimension"] for d in dims_info]
                dimensions_desc = ", ".join(f"{d['dimension']} ({d['description'] or 'no specific details'})" for d in dims_info)
            else:
                dimensions = DEFAULT_DIMENSIONS
                dimensions_desc = ", ".join(DEFAULT_DIMENSIONS)
                
            json_schema_props = ",\n    ".join(f'"{d}": <int 1-5>' for d in dimensions)
            
            prompt = f"""You are Jasmine, a friendly but rigorous technical interviewer. You just finished interviewing {candidate_name}.
Review the transcript below and provide a final holistic grade across {len(dimensions)} dimensions: {dimensions_desc}.
Score each dimension from 1 to 5. Also provide overall_feedback (a few paragraphs of constructive feedback identifying overarching patterns, addressed directly to {candidate_name}).

Output strictly in JSON format matching this schema:
{{
  "rubric_scores": {{
    {json_schema_props}
  }},
  "overall_feedback": "<string>"
}}

TRANSCRIPT:
{transcript}
"""
            completion = client.chat.completions.create(
                model="openai/gpt-oss-20b",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            
            result_json = completion.choices[0].message.content
            print(f"[Grade] Raw Groq response for session {session_id}: {result_json}")
            result = json.loads(result_json)
            
            # Insert scorecard
            cur.execute("""
                INSERT INTO scorecards (session_id, rubric_scores, overall_feedback)
                VALUES (%s, %s, %s)
                RETURNING id
            """, (session_id, json.dumps(result["rubric_scores"]), result["overall_feedback"]))
            
            scorecard_id = cur.fetchone()[0]
            
            return jsonify({
                "scorecard_id": scorecard_id,
                "rubric_scores": result["rubric_scores"],
                "overall_feedback": result["overall_feedback"]
            })



@app.post("/parse-document")
def parse_document():
    """
    Accepts:
    {
      "type": "resume" | "job_description",
      "id": "uuid"
    }
    """
    data = request.json
    doc_type = data.get("type")
    doc_id = data.get("id")

    if doc_type not in ["resume", "job_description"] or not doc_id:
        return jsonify({"error": "Invalid type or id"}), 400

    try:
        with psycopg.connect(DATABASE_URL) as conn:
            with conn.cursor() as cur:
                if doc_type == "resume":
                    cur.execute("SELECT file_path, file_type FROM resumes WHERE id = %s", (doc_id,))
                    row = cur.fetchone()
                    if not row:
                        return jsonify({"error": "Resume not found"}), 404
                    
                    file_path, file_type = row
                    
                    supabase_url = os.getenv("NEXT_PUBLIC_SUPABASE_URL") or os.getenv("SUPABASE_URL")
                    supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
                    
                    if not supabase_url or not supabase_key:
                        return jsonify({"error": "Supabase credentials missing"}), 500
                        
                    from supabase import create_client, Client
                    supabase: Client = create_client(supabase_url, supabase_key)
                    
                    # Download file
                    res = supabase.storage.from_("resumes").download(file_path)
                    
                    extracted_text = ""
                    if file_type == "pdf":
                        import fitz
                        doc = fitz.open(stream=res, filetype="pdf")
                        for page in doc:
                            extracted_text += page.get_text() + "\n"
                    elif file_type == "docx":
                        import docx
                        import io
                        doc = docx.Document(io.BytesIO(res))
                        for para in doc.paragraphs:
                            extracted_text += para.text + "\n"
                    else:
                        return jsonify({"error": "Unsupported file type"}), 400
                        
                    # Strip excessive whitespace
                    import re
                    extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text).strip()
                    
                    # Embed (truncate to 3000 chars)
                    embedding = model.encode(extracted_text[:3000]).tolist()
                    
                    cur.execute("""
                        UPDATE resumes 
                        SET raw_text = %s, embedding = %s, parsed_at = now()
                        WHERE id = %s
                    """, (extracted_text, embedding, doc_id))
                    
                elif doc_type == "job_description":
                    cur.execute("SELECT raw_text FROM job_descriptions WHERE id = %s", (doc_id,))
                    row = cur.fetchone()
                    if not row:
                        return jsonify({"error": "Job description not found"}), 404
                        
                    raw_text = row[0]
                    if not raw_text:
                        return jsonify({"error": "Job description text is empty"}), 400
                        
                    # Embed (truncate to 3000 chars)
                    embedding = model.encode(raw_text[:3000]).tolist()
                    
                    cur.execute("""
                        UPDATE job_descriptions 
                        SET embedding = %s
                        WHERE id = %s
                    """, (embedding, doc_id))

        return jsonify({"success": True})
        
    except Exception as e:
        print(f"Error parsing document: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False)
